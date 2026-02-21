"""
================================================================================
PROJECT 1: RAG DOCUMENT PROCESSING SYSTEM
================================================================================
Making engineers faster at reading documents

This prototype demonstrates the complete pipeline:
1. Document ingestion & parsing
2. Smart chunking with overlap
3. Vector embedding generation
4. Storage in vector database
5. Semantic search & retrieval
6. LLM-powered answer generation with citations

Tech Stack:
- OpenAI (embeddings + GPT-4)
- Pinecone (vector database)
- LangChain (orchestration)
- PyPDF2, python-docx (document parsing)
================================================================================
"""

import os
import hashlib
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import json

# ============================================================================
# CONFIGURATION
# ============================================================================

@dataclass
class RAGConfig:
    """Configuration for the RAG system"""
    # Chunking settings
    chunk_size: int = 800  # tokens per chunk
    chunk_overlap: int = 100  # overlap between chunks
    
    # Embedding settings
    embedding_model: str = "text-embedding-ada-002"
    embedding_dimensions: int = 1536
    
    # Retrieval settings
    top_k: int = 10  # initial retrieval
    rerank_top_k: int = 5  # after reranking
    similarity_threshold: float = 0.75
    
    # LLM settings
    llm_model: str = "gpt-4-turbo-preview"
    max_tokens: int = 1500
    temperature: float = 0.1  # low for factual accuracy


# ============================================================================
# DOCUMENT MODELS
# ============================================================================

@dataclass
class DocumentChunk:
    """A chunk of a document with metadata"""
    id: str
    content: str
    metadata: Dict
    embedding: Optional[List[float]] = None
    
    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "content": self.content,
            "metadata": self.metadata,
            "embedding": self.embedding
        }


@dataclass
class RetrievedContext:
    """Context retrieved from vector search"""
    chunk: DocumentChunk
    score: float
    
    
@dataclass 
class RAGResponse:
    """Response from the RAG system"""
    answer: str
    citations: List[Dict]
    confidence: float
    sources_used: int
    processing_time_ms: int


# ============================================================================
# DOCUMENT PARSER
# ============================================================================

class DocumentParser:
    """
    Parses various document formats into plain text.
    Handles PDFs, Word docs, and plain text.
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.md']
    
    def parse(self, file_path: str) -> Tuple[str, Dict]:
        """
        Parse a document and return (text, metadata)
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif ext in ['.txt', '.md']:
            return self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Parse PDF using PyPDF2"""
        # In production: use PyPDF2 or pdfplumber
        # This is a simplified example
        
        """
        import PyPDF2
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page_num, page in enumerate(reader.pages):
                text += f"\n[Page {page_num + 1}]\n"
                text += page.extract_text()
            
            metadata = {
                "filename": os.path.basename(file_path),
                "pages": len(reader.pages),
                "format": "pdf"
            }
        
        return text, metadata
        """
        
        # Placeholder for demo
        return self._placeholder_parse(file_path, "pdf")
    
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Parse Word document"""
        # In production: use python-docx
        
        """
        from docx import Document
        
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        metadata = {
            "filename": os.path.basename(file_path),
            "paragraphs": len(doc.paragraphs),
            "format": "docx"
        }
        
        return text, metadata
        """
        
        return self._placeholder_parse(file_path, "docx")
    
    def _parse_text(self, file_path: str) -> Tuple[str, Dict]:
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        metadata = {
            "filename": os.path.basename(file_path),
            "characters": len(text),
            "format": "txt"
        }
        
        return text, metadata
    
    def _placeholder_parse(self, file_path: str, format: str) -> Tuple[str, Dict]:
        """Placeholder for demo purposes"""
        return (
            f"[Parsed content from {os.path.basename(file_path)}]",
            {"filename": os.path.basename(file_path), "format": format}
        )


# ============================================================================
# SMART CHUNKER
# ============================================================================

class SmartChunker:
    """
    Intelligently chunks documents for optimal retrieval.
    
    Key strategies:
    - Respect sentence boundaries
    - Respect paragraph boundaries  
    - Maintain overlap for context continuity
    - Preserve tables and lists as units
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.sentence_endings = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
        self.paragraph_marker = '\n\n'
    
    def chunk(self, text: str, metadata: Dict) -> List[DocumentChunk]:
        """
        Split text into chunks with smart boundaries
        """
        chunks = []
        
        # First, split by paragraphs
        paragraphs = text.split(self.paragraph_marker)
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Estimate token count (rough: 1 token ≈ 4 chars)
            estimated_tokens = len(current_chunk + para) // 4
            
            if estimated_tokens > self.config.chunk_size and current_chunk:
                # Save current chunk
                chunk = self._create_chunk(
                    content=current_chunk,
                    metadata=metadata,
                    index=chunk_index
                )
                chunks.append(chunk)
                chunk_index += 1
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap(current_chunk)
                current_chunk = overlap_text + para
            else:
                current_chunk += ("\n\n" if current_chunk else "") + para
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunk = self._create_chunk(
                content=current_chunk,
                metadata=metadata,
                index=chunk_index
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(self, content: str, metadata: Dict, index: int) -> DocumentChunk:
        """Create a DocumentChunk with unique ID"""
        # Generate deterministic ID based on content
        content_hash = hashlib.md5(content.encode()).hexdigest()[:12]
        chunk_id = f"{metadata.get('filename', 'doc')}_{index}_{content_hash}"
        
        chunk_metadata = {
            **metadata,
            "chunk_index": index,
            "char_count": len(content),
            "estimated_tokens": len(content) // 4,
            "created_at": datetime.utcnow().isoformat()
        }
        
        return DocumentChunk(
            id=chunk_id,
            content=content,
            metadata=chunk_metadata
        )
    
    def _get_overlap(self, text: str) -> str:
        """Get the last N characters for overlap"""
        overlap_chars = self.config.chunk_overlap * 4  # rough token to char
        
        if len(text) <= overlap_chars:
            return text
        
        # Try to break at sentence boundary
        overlap_text = text[-overlap_chars:]
        
        for ending in self.sentence_endings:
            idx = overlap_text.find(ending)
            if idx != -1:
                return overlap_text[idx + len(ending):]
        
        return overlap_text


# ============================================================================
# EMBEDDING GENERATOR
# ============================================================================

class EmbeddingGenerator:
    """
    Generates vector embeddings using OpenAI's API.
    Includes batching and caching for efficiency.
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.cache = {}  # Simple in-memory cache
        
    def generate(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings for a list of texts.
        Uses batching for efficiency.
        """
        # In production: use OpenAI API
        
        """
        from openai import OpenAI
        client = OpenAI()
        
        embeddings = []
        batch_size = 100  # OpenAI limit
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            
            response = client.embeddings.create(
                model=self.config.embedding_model,
                input=batch
            )
            
            batch_embeddings = [item.embedding for item in response.data]
            embeddings.extend(batch_embeddings)
        
        return embeddings
        """
        
        # Placeholder: return dummy embeddings
        import random
        return [
            [random.random() for _ in range(self.config.embedding_dimensions)]
            for _ in texts
        ]
    
    def generate_single(self, text: str) -> List[float]:
        """Generate embedding for a single text"""
        # Check cache first
        cache_key = hashlib.md5(text.encode()).hexdigest()
        if cache_key in self.cache:
            return self.cache[cache_key]
        
        embedding = self.generate([text])[0]
        self.cache[cache_key] = embedding
        
        return embedding


# ============================================================================
# VECTOR STORE
# ============================================================================

class VectorStore:
    """
    Vector database interface using Pinecone.
    Handles indexing, search, and metadata filtering.
    """
    
    def __init__(self, config: RAGConfig, index_name: str = "engineering-docs"):
        self.config = config
        self.index_name = index_name
        self._init_index()
    
    def _init_index(self):
        """Initialize Pinecone index"""
        # In production: use Pinecone
        
        """
        import pinecone
        
        pinecone.init(
            api_key=os.environ["PINECONE_API_KEY"],
            environment=os.environ["PINECONE_ENV"]
        )
        
        # Create index if doesn't exist
        if self.index_name not in pinecone.list_indexes():
            pinecone.create_index(
                name=self.index_name,
                dimension=self.config.embedding_dimensions,
                metric="cosine"
            )
        
        self.index = pinecone.Index(self.index_name)
        """
        
        # Placeholder: in-memory storage
        self.vectors = {}
    
    def upsert(self, chunks: List[DocumentChunk]):
        """
        Insert or update chunks in the vector store
        """
        # In production: batch upsert to Pinecone
        
        """
        vectors = [
            {
                "id": chunk.id,
                "values": chunk.embedding,
                "metadata": {
                    "content": chunk.content[:1000],  # Pinecone metadata limit
                    **chunk.metadata
                }
            }
            for chunk in chunks
        ]
        
        # Batch upsert
        batch_size = 100
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            self.index.upsert(vectors=batch)
        """
        
        # Placeholder
        for chunk in chunks:
            self.vectors[chunk.id] = chunk
        
        print(f"Upserted {len(chunks)} chunks to vector store")
    
    def search(
        self, 
        query_embedding: List[float], 
        top_k: int = 10,
        filter: Optional[Dict] = None
    ) -> List[RetrievedContext]:
        """
        Search for similar vectors
        """
        # In production: query Pinecone
        
        """
        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True,
            filter=filter
        )
        
        contexts = []
        for match in results.matches:
            chunk = DocumentChunk(
                id=match.id,
                content=match.metadata.get("content", ""),
                metadata=match.metadata
            )
            contexts.append(RetrievedContext(
                chunk=chunk,
                score=match.score
            ))
        
        return contexts
        """
        
        # Placeholder: return random chunks
        import random
        contexts = []
        for chunk_id, chunk in list(self.vectors.items())[:top_k]:
            contexts.append(RetrievedContext(
                chunk=chunk,
                score=random.uniform(0.7, 0.95)
            ))
        return contexts


# ============================================================================
# RERANKER
# ============================================================================

class Reranker:
    """
    Reranks retrieved results for better precision.
    Uses a cross-encoder model for more accurate relevance scoring.
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
    
    def rerank(
        self, 
        query: str, 
        contexts: List[RetrievedContext]
    ) -> List[RetrievedContext]:
        """
        Rerank contexts based on query relevance
        """
        # In production: use a cross-encoder like sentence-transformers
        
        """
        from sentence_transformers import CrossEncoder
        
        model = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2')
        
        pairs = [(query, ctx.chunk.content) for ctx in contexts]
        scores = model.predict(pairs)
        
        for ctx, score in zip(contexts, scores):
            ctx.score = float(score)
        
        # Sort by new scores
        contexts.sort(key=lambda x: x.score, reverse=True)
        
        return contexts[:self.config.rerank_top_k]
        """
        
        # Placeholder: just return top-k
        contexts.sort(key=lambda x: x.score, reverse=True)
        return contexts[:self.config.rerank_top_k]


# ============================================================================
# ANSWER GENERATOR
# ============================================================================

class AnswerGenerator:
    """
    Generates answers using GPT-4 with retrieved context.
    Ensures citations and prevents hallucination.
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        
        self.system_prompt = """You are a helpful engineering assistant that answers questions based on provided documentation.

CRITICAL RULES:
1. ONLY use information from the provided context
2. ALWAYS cite your sources using [Source N] format
3. If the context doesn't contain the answer, say "I don't have information about that in the provided documents"
4. Be precise and technical when appropriate
5. If information is uncertain, say so

FORMAT:
- Provide a clear, direct answer
- Include relevant technical details
- Cite every factual claim with [Source N]
- End with a confidence assessment"""

    def generate(
        self, 
        query: str, 
        contexts: List[RetrievedContext]
    ) -> RAGResponse:
        """
        Generate an answer with citations
        """
        import time
        start_time = time.time()
        
        # Build context string with source numbers
        context_parts = []
        citations = []
        
        for i, ctx in enumerate(contexts, 1):
            context_parts.append(f"[Source {i}]\n{ctx.chunk.content}")
            citations.append({
                "source_number": i,
                "document": ctx.chunk.metadata.get("filename", "Unknown"),
                "chunk_id": ctx.chunk.id,
                "relevance_score": round(ctx.score, 3)
            })
        
        context_string = "\n\n---\n\n".join(context_parts)
        
        # Build the prompt
        user_prompt = f"""Context from documents:

{context_string}

---

Question: {query}

Please provide a comprehensive answer based on the context above. Cite your sources."""

        # In production: call OpenAI API
        
        """
        from openai import OpenAI
        client = OpenAI()
        
        response = client.chat.completions.create(
            model=self.config.llm_model,
            messages=[
                {"role": "system", "content": self.system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=self.config.max_tokens,
            temperature=self.config.temperature
        )
        
        answer = response.choices[0].message.content
        """
        
        # Placeholder answer
        answer = f"""Based on the provided documentation, here is the answer to your question about "{query}":

The relevant information indicates that [technical details would be extracted from the context here]. [Source 1]

Additionally, the documentation specifies [more specific information]. [Source 2]

Key points to note:
- Point 1 from the documentation [Source 1]
- Point 2 with technical specifications [Source 3]

Confidence: High - Multiple sources confirm this information."""

        # Calculate confidence based on retrieval scores
        avg_score = sum(ctx.score for ctx in contexts) / len(contexts) if contexts else 0
        confidence = min(avg_score, 0.95)  # Cap at 95%
        
        processing_time = int((time.time() - start_time) * 1000)
        
        return RAGResponse(
            answer=answer,
            citations=citations,
            confidence=round(confidence, 2),
            sources_used=len(contexts),
            processing_time_ms=processing_time
        )


# ============================================================================
# MAIN RAG PIPELINE
# ============================================================================

class RAGPipeline:
    """
    Complete RAG pipeline that orchestrates all components.
    """
    
    def __init__(self, config: Optional[RAGConfig] = None):
        self.config = config or RAGConfig()
        
        # Initialize components
        self.parser = DocumentParser()
        self.chunker = SmartChunker(self.config)
        self.embedder = EmbeddingGenerator(self.config)
        self.vector_store = VectorStore(self.config)
        self.reranker = Reranker(self.config)
        self.generator = AnswerGenerator(self.config)
        
        print("RAG Pipeline initialized")
        print(f"  - Chunk size: {self.config.chunk_size} tokens")
        print(f"  - Embedding model: {self.config.embedding_model}")
        print(f"  - LLM model: {self.config.llm_model}")
    
    # -------------------------------------------------------------------------
    # INDEXING PHASE
    # -------------------------------------------------------------------------
    
    def index_document(self, file_path: str) -> int:
        """
        Index a single document into the vector store.
        Returns the number of chunks created.
        """
        print(f"\n📄 Indexing: {file_path}")
        
        # Step 1: Parse document
        print("  → Parsing document...")
        text, metadata = self.parser.parse(file_path)
        
        # Step 2: Chunk document
        print("  → Chunking...")
        chunks = self.chunker.chunk(text, metadata)
        print(f"    Created {len(chunks)} chunks")
        
        # Step 3: Generate embeddings
        print("  → Generating embeddings...")
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedder.generate(texts)
        
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
        
        # Step 4: Store in vector database
        print("  → Storing in vector database...")
        self.vector_store.upsert(chunks)
        
        print(f"  ✅ Indexed {len(chunks)} chunks")
        return len(chunks)
    
    def index_directory(self, dir_path: str) -> Dict:
        """
        Index all documents in a directory.
        Returns statistics about the indexing.
        """
        stats = {
            "documents_processed": 0,
            "total_chunks": 0,
            "errors": []
        }
        
        for filename in os.listdir(dir_path):
            file_path = os.path.join(dir_path, filename)
            
            if not os.path.isfile(file_path):
                continue
            
            ext = os.path.splitext(filename)[1].lower()
            if ext not in self.parser.supported_formats:
                continue
            
            try:
                chunks = self.index_document(file_path)
                stats["documents_processed"] += 1
                stats["total_chunks"] += chunks
            except Exception as e:
                stats["errors"].append({
                    "file": filename,
                    "error": str(e)
                })
        
        return stats
    
    # -------------------------------------------------------------------------
    # QUERY PHASE
    # -------------------------------------------------------------------------
    
    def query(
        self, 
        question: str,
        filter: Optional[Dict] = None
    ) -> RAGResponse:
        """
        Answer a question using the RAG pipeline.
        """
        print(f"\n❓ Query: {question}")
        
        # Step 1: Embed the query
        print("  → Embedding query...")
        query_embedding = self.embedder.generate_single(question)
        
        # Step 2: Retrieve relevant chunks
        print(f"  → Searching vector store (top {self.config.top_k})...")
        contexts = self.vector_store.search(
            query_embedding=query_embedding,
            top_k=self.config.top_k,
            filter=filter
        )
        print(f"    Found {len(contexts)} relevant chunks")
        
        # Step 3: Rerank for precision
        print(f"  → Reranking to top {self.config.rerank_top_k}...")
        contexts = self.reranker.rerank(question, contexts)
        
        # Step 4: Generate answer
        print("  → Generating answer...")
        response = self.generator.generate(question, contexts)
        
        print(f"  ✅ Answer generated (confidence: {response.confidence})")
        
        return response


# ============================================================================
# USAGE EXAMPLE
# ============================================================================

def main():
    """Demonstrate the RAG pipeline"""
    
    print("=" * 70)
    print("RAG DOCUMENT PROCESSING SYSTEM - DEMO")
    print("=" * 70)
    
    # Initialize pipeline
    config = RAGConfig(
        chunk_size=500,
        chunk_overlap=50,
        top_k=10,
        rerank_top_k=5
    )
    
    pipeline = RAGPipeline(config)
    
    # Simulate indexing documents
    print("\n" + "=" * 70)
    print("PHASE 1: INDEXING DOCUMENTS")
    print("=" * 70)
    
    # In production, you would index real documents:
    # pipeline.index_directory("/path/to/engineering/docs")
    
    # For demo, let's simulate some chunks
    demo_chunks = [
        DocumentChunk(
            id="pump-maintenance-001",
            content="""Pump Maintenance Procedure for Unit 7
            
            1. Isolate the pump using lockout-tagout (LOTO) procedure
            2. Verify zero energy state
            3. Remove coupling guard
            4. Check impeller condition - replace if blade thickness < 3mm
            5. Inspect mechanical seal for wear
            6. Check bearing temperature logs from past 30 days
            7. Lubricate as per schedule in Appendix B""",
            metadata={
                "filename": "Unit7-Maintenance-SOP.pdf",
                "page": 23,
                "section": "Pump Maintenance"
            }
        ),
        DocumentChunk(
            id="safety-loto-001", 
            content="""Lockout-Tagout (LOTO) Safety Procedure
            
            Before any maintenance work:
            1. Notify control room of planned maintenance
            2. Identify all energy sources (electrical, hydraulic, pneumatic)
            3. Shut down equipment using normal procedures
            4. Apply locks and tags at each isolation point
            5. Verify zero energy state with appropriate testing
            6. Document all locks in the LOTO log""",
            metadata={
                "filename": "Safety-Manual.pdf",
                "page": 45,
                "section": "LOTO Procedures"
            }
        )
    ]
    
    # Generate embeddings and store
    for chunk in demo_chunks:
        chunk.embedding = pipeline.embedder.generate_single(chunk.content)
    
    pipeline.vector_store.upsert(demo_chunks)
    
    # Query the system
    print("\n" + "=" * 70)
    print("PHASE 2: QUERYING")
    print("=" * 70)
    
    questions = [
        "What is the procedure for pump maintenance on Unit 7?",
        "How do I perform lockout-tagout?",
        "When should I replace the pump impeller?"
    ]
    
    for question in questions:
        response = pipeline.query(question)
        
        print("\n" + "-" * 50)
        print(f"Q: {question}")
        print("-" * 50)
        print(f"\nA: {response.answer}")
        print(f"\n📚 Sources cited: {response.sources_used}")
        print(f"🎯 Confidence: {response.confidence}")
        print(f"⏱️  Processing time: {response.processing_time_ms}ms")
        
        print("\n📎 Citations:")
        for citation in response.citations:
            print(f"   [{citation['source_number']}] {citation['document']} "
                  f"(relevance: {citation['relevance_score']})")


if __name__ == "__main__":
    main()
